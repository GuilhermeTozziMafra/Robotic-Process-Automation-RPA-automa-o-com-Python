from docx import Document

#Aumenta o tamanho da da letra da fonte
from docx.shared import Pt

#Muda a cor do texto
from docx.shared import RGBColor

from openpyxl import load_workbook

#Pegando o caminho do arquivo
nome_arquivo_alunos = "DadosAlunos.xlsx"
planilhaDadosAlunos = load_workbook(nome_arquivo_alunos)

#Selecionando a sheet/planilha/aba
sheet_selecionada = planilhaDadosAlunos["Nomes"]

#for = para
for linha in range(2, len(sheet_selecionada["A"]) + 1):

    #Abre o arquivo do word
    arquivoWord = Document("Certificado3.docx")

    #Seleciona o estilo
    estilo = arquivoWord.styles["Normal"]

    #Pegando o nome do aluno quando passar na célula
    nomeAluno = sheet_selecionada['A%s' % linha].value
    dia = sheet_selecionada['B%s' % linha].value
    mes = sheet_selecionada['C%s' % linha].value
    ano = sheet_selecionada['D%s' % linha].value
    nomeCurso = sheet_selecionada['E%s' % linha].value
    nomeInstrutor = sheet_selecionada['F%s' % linha].value

    #for = para
    for paragrafo in arquivoWord.paragraphs:
        #if = se
        if "@nome" in paragrafo.text:
            paragrafo.text = nomeAluno
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

        paragrafoP1 = "Concluiu com sucesso o curso de "
        paragrafoP2 = ", com a carga horária de 20 horas, promovido pela escola de Cursos Online em "
        terceiraParteParagrafo = f"{paragrafoP2} {dia} de {mes} de {ano}."

        if "escola" in paragrafo.text:
            paragrafo.text = paragrafoP1
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)
            adicionaNovaPalavra = paragrafo.add_run(nomeCurso)
            adicionaNovaPalavra.font.color.rgb = RGBColor(255, 0, 0) #Muda para cor vermelho
            adicionaNovaPalavra.underline = True #Sublinhado
            adicionaNovaPalavra.bold = True #Negrito
            adicionaNovaPalavra = paragrafo.add_run(terceiraParteParagrafo)
            adicionaNovaPalavra.font.color.rgb = RGBColor(0, 0, 0)  # Muda para cor preto


        if "Instrutor" in paragrafo.text:
            paragrafo.text = nomeInstrutor + " - Instrutor"
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

    #Pegando o caminho da pasta e configurando o nome do certificado
    caminhoCertificados = "A:\\Curso Lógica de Programação\\Python\\Word\\Certificados\\" + nomeAluno + ".docx"

    #Salvando o certificado do aluno
    arquivoWord.save(caminhoCertificados)

print("Certificados gerados com sucesso!")