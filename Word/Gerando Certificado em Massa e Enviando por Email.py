from docx import Document

#Aumenta o tamanho da da letra da fonte
from docx.shared import Pt

#Muda a cor do texto
from docx.shared import RGBColor

from openpyxl import load_workbook

#pip install PyWin32
import win32com.client as win32
outlook = win32.Dispatch("outlook.application")

#Pegando o caminho do arquivo
nome_arquivo_alunos = "DadosAlunosEmail.xlsx"
planilhaDadosAlunos = load_workbook(nome_arquivo_alunos)

#Selecionando a sheet/planilha/aba
sheet_selecionada = planilhaDadosAlunos["Nomes"]

#for = para
for linha in range(2, len(sheet_selecionada["A"]) + 1):

    #Abre o arquivo do word
    arquivoWord = Document("Certificado3.docx")

    #Seleciona o estilo
    estilo = arquivoWord.styles["Normal"]

    #Pegando o nome do aluno quando passar na célula
    nomeAluno = sheet_selecionada['A%s' % linha].value
    dia = sheet_selecionada['B%s' % linha].value
    mes = sheet_selecionada['C%s' % linha].value
    ano = sheet_selecionada['D%s' % linha].value
    nomeCurso = sheet_selecionada['E%s' % linha].value
    nomeInstrutor = sheet_selecionada['F%s' % linha].value
    emailAluno = sheet_selecionada['G%s' % linha].value

    #for = para
    for paragrafo in arquivoWord.paragraphs:
        #if = se
        if "@nome" in paragrafo.text:
            paragrafo.text = nomeAluno
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

        paragrafoP1 = "Concluiu com sucesso o curso de "
        paragrafoP2 = ", com a carga horária de 20 horas, promovido pela escola de Cursos Online em "
        terceiraParteParagrafo = f"{paragrafoP2} {dia} de {mes} de {ano}."

        if "escola" in paragrafo.text:
            paragrafo.text = paragrafoP1
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)
            adicionaNovaPalavra = paragrafo.add_run(nomeCurso)
            adicionaNovaPalavra.font.color.rgb = RGBColor(255, 0, 0) #Muda para cor vermelho
            adicionaNovaPalavra.underline = True #Sublinhado
            adicionaNovaPalavra.bold = True #Negrito
            adicionaNovaPalavra = paragrafo.add_run(terceiraParteParagrafo)
            adicionaNovaPalavra.font.color.rgb = RGBColor(0, 0, 0)  # Muda para cor preto


        if "Instrutor" in paragrafo.text:
            paragrafo.text = nomeInstrutor + " - Instrutor"
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

    #Pegando o caminho da pasta e configurando o nome do certificado
    caminhoCertificados = "A:\\Curso Lógica de Programação\\Python\\Word\\Certificados\\" + nomeAluno + ".docx"

    #Salvando o certificado do aluno
    arquivoWord.save(caminhoCertificados)

    #Pega o primeiro nome do aluno
    primeiroNome = nomeAluno.split(None, 1)[0]

    emailOutlook = outlook.CreateItem(0)
    emailOutlook.To = emailAluno
    emailOutlook.Subject = "Certificado " + nomeAluno
    emailOutlook.HTMLBody = f"""
        <p>Boa noite {primeiroNome}.</p>
        <p>Segue seu <b>certificado.</b></p>
        <p>Atenciosamente,</p>
        <p><img src="C:\Users\guilherme.mafra\Desktop\Python\Word.></p>
    """
    #Adiciona anexo / certificado no email
    emailOutlook.Attachments.Add(caminhoCertificados)
    #.save() = Salvar como Draft no outlook
    #.send ou .send() = Enviar o email
    emailOutlook.save()

print("Certificados gerados com sucesso!")